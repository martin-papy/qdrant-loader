"""Behaviour tests for the XLSX post-processing of the docling conversion path.

Two defects of docling's Excel backend, fixed once on the structured
``DoclingDocument`` so BOTH the markdown export and the HybridChunker benefit:

* Defect 1 — merged-cell duplication. docling stores a merged cell once in
  ``table.data.table_cells`` with span info, but the derived ``grid`` (and the
  markdown/chunk serializers that read it) repeat the cell's text into every
  spanned position. A title merged across N columns shows up N times.
* Defect 2 — sheet names dropped. The sheet names survive as
  ``doc.groups[*].name == "sheet: <name>"`` but neither the markdown export nor
  the chunker surfaces them. markitdown emitted ``## Sheet: <name>`` headings.

These build a deterministic workbook with openpyxl (two named sheets, one table
with a cell merged across several columns and rows) and convert it through the
real ``ConversionService(FileConversionConfig(engine="docling"))`` — no mocks.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from qdrant_loader.core.file_conversion.conversion_config import FileConversionConfig

# A long, distinctive sentence so a single merged anchor is unambiguous to count.
MERGED_TITLE = "ACME GLOBAL VENDOR RFP MASTER TITLE BANNER SENTENCE"
SHEET_ONE = "A. Vendor Capacity"
SHEET_TWO = "B. Vendor Compliance"


@pytest.fixture
def merged_cell_workbook(tmp_path: Path) -> Path:
    """A 2-sheet xlsx with a cell merged across 5 columns x 2 rows on sheet one.

    Sheet one: row 1-2 / col A-E merged carrying ``MERGED_TITLE``; below it a
    couple of ordinary data rows. Sheet two: a small ordinary table. The merge
    is what triggers docling's grid duplication; the two distinct sheet names
    are what the heading fix must surface.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = SHEET_ONE
    ws1.merge_cells(start_row=1, start_column=1, end_row=2, end_column=5)
    ws1.cell(row=1, column=1, value=MERGED_TITLE)
    ws1.cell(row=3, column=1, value="Capability")
    ws1.cell(row=3, column=2, value="Score")
    ws1.cell(row=4, column=1, value="Throughput")
    ws1.cell(row=4, column=2, value="High")

    ws2 = wb.create_sheet(title=SHEET_TWO)
    ws2.cell(row=1, column=1, value="Control")
    ws2.cell(row=1, column=2, value="Status")
    ws2.cell(row=2, column=1, value="Encryption")
    ws2.cell(row=2, column=2, value="Met")

    path = tmp_path / "merged_two_sheets.xlsx"
    wb.save(path)
    return path


def _convert(path: Path):
    service = ConversionServiceFactory()
    return service.convert(str(path))


def ConversionServiceFactory():
    from qdrant_loader.core.conversion.service import ConversionService

    return ConversionService(FileConversionConfig(engine="docling"))


# ── Defect 1: merged-cell text appears exactly once ──────────────────────────


def test_merged_cell_text_appears_once_in_markdown(merged_cell_workbook):
    """A cell merged across 5x2 must contribute its text once, not 10 times."""
    converted = _convert(merged_cell_workbook)
    markdown = converted.converted_document.to_markdown()

    assert markdown.count(MERGED_TITLE) == 1, (
        f"merged title duplicated: appeared {markdown.count(MERGED_TITLE)}x"
    )


def test_table_cells_retain_span_metadata(merged_cell_workbook):
    """The fix must NOT flatten spans in table_cells — only stop the serialized
    duplication. The anchor cell keeps its multi-row/col span; empty fillers
    occupy the vacated positions."""
    converted = _convert(merged_cell_workbook)
    doc = converted.converted_document.document

    spanned = [
        cell
        for table in doc.tables
        for cell in table.data.table_cells
        if (cell.col_span > 1 or cell.row_span > 1) and MERGED_TITLE in cell.text
    ]
    assert spanned, "the merged anchor cell must still carry its span metadata"
    anchor = spanned[0]
    assert anchor.col_span >= 2 and anchor.row_span >= 2


# ── Defect 2: each sheet name appears as a markdown heading ───────────────────


def test_each_sheet_name_is_a_markdown_heading(merged_cell_workbook):
    converted = _convert(merged_cell_workbook)
    markdown = converted.converted_document.to_markdown()

    heading_lines = {
        line.strip() for line in markdown.splitlines() if line.lstrip().startswith("#")
    }
    assert any(SHEET_ONE in line for line in heading_lines), (
        f"sheet one heading missing; headings were {heading_lines}"
    )
    assert any(SHEET_TWO in line for line in heading_lines), (
        f"sheet two heading missing; headings were {heading_lines}"
    )


def test_headings_use_sheet_prefix(merged_cell_workbook):
    """Mirror markitdown's ``## Sheet: <name>`` convention."""
    converted = _convert(merged_cell_workbook)
    markdown = converted.converted_document.to_markdown()

    assert "Sheet: " + SHEET_ONE in markdown
    assert "Sheet: " + SHEET_TWO in markdown


# ── Defect 2 must NOT touch non-spreadsheet inputs ────────────────────────────


def test_docx_is_not_given_sheet_headings(tmp_path):
    """Only XLSX inputs get sheet headers; a prose docx is left alone."""
    from docx import Document as DocxDocument

    docx_path = tmp_path / "prose.docx"
    d = DocxDocument()
    d.add_heading("Real Heading", level=1)
    d.add_paragraph("Some lorem ipsum prose body text for the conversion.")
    d.save(docx_path)

    converted = _convert(docx_path)
    markdown = converted.converted_document.to_markdown()

    assert "Sheet:" not in markdown


# ── Chunker path benefits from the same fix ───────────────────────────────────


def test_chunks_do_not_duplicate_merged_text_and_carry_sheet_context(
    merged_cell_workbook, test_settings
):
    """The chunker consumes the same post-processed DoclingDocument, so its
    chunks must (a) never repeat the merged title within a single chunk and
    (b) carry the originating sheet name in their structure heading_path."""
    from qdrant_loader.connectors.localfile import LocalFileConnector  # noqa: F401
    from qdrant_loader.core.chunking.strategy.docling_strategy import (
        DoclingChunkingStrategy,
    )
    from qdrant_loader.core.document import Document

    converted = _convert(merged_cell_workbook)
    document = Document(
        title="merged_two_sheets.xlsx",
        content=converted.content,
        content_type="xlsx",
        source="test",
        source_type="localfile",
        url="file://merged_two_sheets.xlsx",
        metadata={"conversion_method": "docling"},
    )
    document.converted_document = converted.converted_document

    strategy = DoclingChunkingStrategy(test_settings)
    chunks = strategy.chunk_document(document)

    assert chunks, "expected at least one chunk"
    for chunk in chunks:
        assert chunk.content.count(MERGED_TITLE) <= 1, (
            "merged title duplicated within a chunk"
        )

    heading_blobs = [
        " ".join(chunk.metadata.get("structure", {}).get("heading_path", []))
        for chunk in chunks
    ]
    joined = " ".join(heading_blobs)
    assert SHEET_ONE in joined or SHEET_TWO in joined, (
        f"no chunk carried a sheet name in heading_path; got {heading_blobs}"
    )
